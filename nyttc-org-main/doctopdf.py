import docx
import pdfminer
from io import StringIO

def convert_pdf_to_word(path):
    with open(path, 'rb') as file:
        # Extract text from PDF
        text = extract_text_from_pdf(file)
        
        # Create docx document
        document = docx.Document()
        
        # Add text to document
        document.add_paragraph(text)
        
        # Save document
        document.save(path + '.docx')

def extract_text_from_pdf(file):
    # Create a PDF parser object
    parser = pdfminer.PDFParser(file)
    
    # Create a PDF document object
    doc = pdfminer.PDFDocument(parser)
    
    # Check if the document allows text extraction
    if not doc.is_extractable:
        raise Exception('Text extraction is not allowed')
    
    # Create a PDF resource manager object
    rsrcmgr = pdfminer.PDFResourceManager()
    
    # Create a StringIO object for storing the extracted text
    string_io = StringIO()
    
    # Create a PDF device object
    device = pdfminer.TextConverter(rsrcmgr, string_io)
    
    # Create a PDF interpreter object
    interpreter = pdfminer.PDFPageInterpreter(rsrcmgr, device)
    
    # Extract text from each page and store it in the StringIO object
    for page in doc.get_pages():
        interpreter.process_page(page)
        
    # Get the text from the StringIO object
    text = string_io.getvalue()
    
    # Cleanup
    string_io.close()
    device.close()
    
    # Return the extracted text
    return text

# Example usage
convert_pdf_to_word('document.pdf')
